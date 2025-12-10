from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from core.models.data_classes import Settings, Figure, Table, Citation
from core.utils.helpers import format_citation_apa
from typing import List
import os
import re
import traceback
import matplotlib.pyplot as plt
import io
import latex2mathml.converter
from lxml import etree

# XSLT to convert MathML to OMML (Word's equation format)
MATHML_TO_OMML_XSLT = None

def get_mathml_to_omml_xslt():
    """Load the MathML to OMML XSLT stylesheet."""
    global MATHML_TO_OMML_XSLT
    if MATHML_TO_OMML_XSLT is None:
        # This XSLT is bundled with Microsoft Office
        # On Mac, it's typically at this location
        xslt_paths = [
            "/Applications/Microsoft Word.app/Contents/Resources/MML2OMML.XSL",
            "/Applications/Microsoft Office/Microsoft Word.app/Contents/Resources/MML2OMML.XSL",
            # Fallback for different Office versions
            os.path.expanduser("~/Applications/Microsoft Word.app/Contents/Resources/MML2OMML.XSL"),
        ]
        
        for xslt_path in xslt_paths:
            if os.path.exists(xslt_path):
                xslt_doc = etree.parse(xslt_path)
                MATHML_TO_OMML_XSLT = etree.XSLT(xslt_doc)
                print(f"Loaded MML2OMML.XSL from: {xslt_path}")
                break
        
        if MATHML_TO_OMML_XSLT is None:
            print("Warning: MML2OMML.XSL not found. LaTeX equations will use fallback image rendering.")
    
    return MATHML_TO_OMML_XSLT

def latex_to_omml(latex_str):
    """Convert LaTeX string to Word OMML (Office Math Markup Language)."""
    try:
        # Clean up the LaTeX string
        clean_latex = latex_str.strip()
        if clean_latex.startswith('$') and clean_latex.endswith('$'):
            clean_latex = clean_latex[1:-1]
        if clean_latex.startswith('$$') and clean_latex.endswith('$$'):
            clean_latex = clean_latex[2:-2]
        
        # Convert LaTeX to MathML
        mathml_str = latex2mathml.converter.convert(clean_latex)
        
        # Get the XSLT transformer
        xslt = get_mathml_to_omml_xslt()
        if xslt is None:
            return None
        
        # Parse MathML and transform to OMML
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
        omml_tree = xslt(mathml_tree)
        
        # Get the oMath element
        omml_element = omml_tree.getroot()
        
        return omml_element
        
    except Exception as e:
        print(f"Error converting LaTeX to OMML: {e}")
        traceback.print_exc()
        return None

def insert_omml_equation(paragraph, latex_str):
    """Insert a LaTeX equation as native Word OMML into a paragraph."""
    omml = latex_to_omml(latex_str)
    if omml is not None:
        # Create a run and append the OMML
        run = paragraph.add_run()
        run._r.append(omml)
        return True
    return False

def render_latex_to_image(latex_str, font_size_pt=12, is_display=False):
    """Renders LaTeX string to an image stream using matplotlib."""
    try:
        # Configure Matplotlib to use STIX (Times-like) for Math
        plt.rcParams['mathtext.fontset'] = 'stix'
        plt.rcParams['font.family'] = 'Times New Roman'
        
        # High DPI for quality
        dpi = 600
        
        # Create figure
        fig = plt.figure(figsize=(8, 2), dpi=dpi)
        fig.patch.set_alpha(0)
        
        # Clean up LaTeX
        render_str = latex_str.strip()
        if render_str.startswith('$$') and render_str.endswith('$$'):
            render_str = render_str[2:-2]
        elif render_str.startswith('$') and render_str.endswith('$'):
            render_str = render_str[1:-1]
        
        # Check for unsupported LaTeX environments - return None to trigger fallback
        unsupported_patterns = [
            r'\\begin\{', r'\\end\{',  # environments like cases, matrix, etc.
            r'\\underbrace', r'\\overbrace',  # braces
            r'\\xrightarrow', r'\\xleftarrow',  # extensible arrows
            r'\\substack',  # stacked subscripts
            r'\\overset', r'\\underset',  # over/under set
        ]
        for pattern in unsupported_patterns:
            if re.search(pattern, render_str):
                print(f"Unsupported LaTeX pattern detected: {pattern} in '{render_str[:50]}...'")
                plt.close(fig)
                return None, 0, 0, 0
            
        # Fix common LaTeX commands
        render_str = re.sub(r'\\displaystyle\s*', '', render_str)
        render_str = re.sub(r'\\textstyle\s*', '', render_str)
        render_str = re.sub(r'\\ge(?![a-zA-Z])', r'\\geq', render_str)
        render_str = re.sub(r'\\le(?![a-zA-Z])', r'\\leq', render_str)
        render_str = re.sub(r'\\text\{', r'\\mathrm{', render_str)
        render_str = render_str.replace(r'\{', r'\lbrace ')
        render_str = render_str.replace(r'\}', r'\rbrace ')
        render_str = re.sub(r'\\arg(?![a-zA-Z])', r'\\mathrm{arg}', render_str)
        render_str = re.sub(r'\\max(?![a-zA-Z])', r'\\mathrm{max}', render_str)
        render_str = re.sub(r'\\min(?![a-zA-Z])', r'\\mathrm{min}', render_str)
        
        render_str = f"${render_str}$"
            
        # Adjust font size to match document text
        effective_font_size = font_size_pt * 0.65  # Inline math smaller to match text height
        if is_display:
            effective_font_size = font_size_pt * 0.60  # Display math also smaller
            
        # Draw text with baseline alignment for accurate descent calculation
        text = fig.text(0.5, 0.5, render_str, fontsize=effective_font_size, ha='center', va='baseline')
        
        # Get bounding box
        renderer = fig.canvas.get_renderer()
        bbox = text.get_window_extent(renderer=renderer)
        
        # Calculate dimensions
        width_in = bbox.width / dpi
        height_in = bbox.height / dpi
        
        # Calculate descent (distance from baseline to bottom)
        # Baseline is at 0.5 * fig_height_in_pixels
        fig_height_px = fig.get_window_extent().height
        baseline_y = 0.5 * fig_height_px
        descent_px = baseline_y - bbox.y0
        descent_in = descent_px / dpi
        
        # Save with tight bounding box to eliminate whitespace
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight', pad_inches=0.01, transparent=True)
        buf.seek(0)
        plt.close(fig)
        
        return buf, height_in, width_in, descent_in
        
    except Exception as e:
        print(f"Error rendering LaTeX: {e}")
        traceback.print_exc()
        try: plt.close(fig)
        except: pass
        return None, 0, 0, 0

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_toc_field(paragraph, field_code):
    """Insert a Word field (like TOC) into a paragraph"""
    run = paragraph.add_run()
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'begin')
    run._element.append(fldChar)

    run = paragraph.add_run()
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = field_code
    run._element.append(instrText)

    run = paragraph.add_run()
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'separate')
    run._element.append(fldChar)

    run = paragraph.add_run()
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'end')
    run._element.append(fldChar)

def set_font_complex(font, font_name, font_size=None, bold=False, italic=False, color=None):
    """Set font properties comprehensively for all script types"""
    font.name = font_name
    if font_size:
        font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color
        
    # Access the underlying XML element to set fonts for all regions
    rPr = font._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def set_run_position(run, pt_val):
    """
    Set the vertical position of a run (raise/lower).
    pt_val: Value in points. Positive raises, negative lowers.
    Word w:position uses half-points (1/144 inch).
    """
    if pt_val == 0: return
    val = int(pt_val * 2)
    rPr = run._element.get_or_add_rPr()
    # w:position element
    position = OxmlElement('w:position')
    position.set(qn('w:val'), str(val))
    rPr.append(position)

def setup_styles(doc, settings: Settings):
    """Configure document styles to match thesis requirements"""
    styles = doc.styles
    
    # --- STYLES ---
    # --- STYLES ---
    style = doc.styles['Normal']
    # Use set_font_complex to ensure font is applied to all script types (important for Vietnamese)
    set_font_complex(style.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))

    # Paragraph formatting for Normal style
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = settings.line_spacing
    # Áp dụng thụt đầu dòng từ settings
    pf.first_line_indent = Cm(settings.indent)
    # Áp dụng khoảng cách đoạn
    pf.space_after = Pt(8.5) 
    pf.space_before = Pt(0)

    # Set Paper Size to A4 (Critical for correct margins)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(settings.margin_top)
        section.bottom_margin = Cm(settings.margin_bottom)
        section.left_margin = Cm(settings.margin_left)
        section.right_margin = Cm(settings.margin_right)
    
    # Heading 1
    style_h1 = styles['Heading 1']
    set_font_complex(style_h1.font, settings.font_family, settings.h1_size, bold=True, color=RGBColor(0, 0, 0))
    paragraph_format = style_h1.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_before = Pt(24)
    paragraph_format.space_after = Pt(18)
    
    # Heading 2
    style_h2 = styles['Heading 2']
    set_font_complex(style_h2.font, settings.font_family, settings.h2_size, bold=True, color=RGBColor(0, 0, 0))
    paragraph_format = style_h2.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format.space_before = Pt(18)
    paragraph_format.space_after = Pt(12)
    
    # Heading 3
    style_h3 = styles['Heading 3']
    set_font_complex(style_h3.font, settings.font_family, settings.h3_size, bold=True, italic=True, color=RGBColor(0, 0, 0))
    paragraph_format = style_h3.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(6)
    
    # Heading 4
    style_h4 = styles['Heading 4']
    set_font_complex(style_h4.font, settings.font_family, settings.font_size, italic=True, color=RGBColor(0, 0, 0))
    paragraph_format = style_h4.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    
    # Heading 5
    try:
        style_h5 = styles.add_style('Heading 5', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style_h5 = styles['Heading 5']
    style_h5.base_style = styles['Normal']
    set_font_complex(style_h5.font, settings.font_family, settings.font_size, italic=True, color=RGBColor(0, 0, 0))
    paragraph_format = style_h5.paragraph_format
    paragraph_format.left_indent = Cm(0.63) # Indent slightly
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    
    # Front Heading (for TOC, List of Tables, etc. - Looks like H1 but not in TOC)
    try:
        style_front = styles.add_style('Front Heading', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style_front = styles['Front Heading']
    style_front.base_style = styles['Normal']
    set_font_complex(style_front.font, settings.font_family, settings.h1_size, bold=True, color=RGBColor(0, 0, 0))
    style_front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_front.paragraph_format.space_before = Pt(24)
    style_front.paragraph_format.space_after = Pt(18)

    # Figure Caption Style
    try:
        style_fig_cap = styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style_fig_cap = styles['Figure Caption']
    style_fig_cap.base_style = styles['Normal']
    set_font_complex(style_fig_cap.font, settings.font_family, settings.font_size, italic=True, color=RGBColor(0, 0, 0))
    style_fig_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_fig_cap.paragraph_format.space_before = Pt(6)
    style_fig_cap.paragraph_format.space_after = Pt(12)

    # Table Caption Style
    try:
        style_tbl_cap = styles.add_style('Table Caption', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style_tbl_cap = styles['Table Caption']
    style_tbl_cap.base_style = styles['Normal']
    set_font_complex(style_tbl_cap.font, settings.font_family, settings.font_size, bold=True, color=RGBColor(0, 0, 0))
    style_tbl_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_tbl_cap.paragraph_format.space_before = Pt(12)
    style_tbl_cap.paragraph_format.space_after = Pt(6)
    
    # List Bullet
    try:
        style_bullet = styles['List Bullet']
        set_font_complex(style_bullet.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
    except:
        pass

def export_to_docx(file_path: str, text: str, settings: Settings, 
                   figures: List[Figure], tables: List[Table], citations: List[Citation],
                   abbreviations: List[dict] = None):
    try:
        doc = Document()
        setup_styles(doc, settings)
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(settings.margin_top)
            section.bottom_margin = Cm(settings.margin_bottom)
            section.left_margin = Cm(settings.margin_left)
            section.right_margin = Cm(settings.margin_right)
        
        # --- FRONT MATTER ---
        
        # 1. MỤC LỤC
        # Add instruction to update fields
        instr_p = doc.add_paragraph("Lưu ý: Nhấn Ctrl+A rồi nhấn F9 (hoặc chuột phải chọn 'Update Field') để cập nhật Mục lục và Danh mục.")
        instr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_complex(instr_p.runs[0].font, settings.font_family, 11, italic=True, color=RGBColor(255, 0, 0))

        toc_p = doc.add_paragraph("MỤC LỤC", style='Front Heading')
        # Insert TOC Field: \o "1-3" includes Heading 1-3, \h hyperlinks, \z hide page numbers in web, \u outline levels
        p = doc.add_paragraph()
        add_toc_field(p, r'TOC \o "1-3" \h \z \u')
        doc.add_page_break()
        
        # 2. DANH MỤC HÌNH ẢNH
        if figures:
            doc.add_paragraph("DANH MỤC HÌNH ẢNH", style='Front Heading')
            # Insert TOC for Figure Caption style
            p = doc.add_paragraph()
            add_toc_field(p, r'TOC \h \z \t "Figure Caption,1"')
            doc.add_page_break()
            
        # 3. DANH MỤC BẢNG BIỂU
        # Note: We don't have a tables list passed in explicitly with captions usually, 
        # but we detect them in text. However, for the TOC to work, we just need the captions in text to use the style.
        # We'll assume if there are tables in text, we want this list. 
        # Since we don't track tables list perfectly in backend yet, we'll just add it if text contains "Bảng"
        if "Bảng" in text:
            doc.add_paragraph("DANH MỤC BẢNG BIỂU", style='Front Heading')
            p = doc.add_paragraph()
            add_toc_field(p, r'TOC \h \z \t "Table Caption,1"')
            doc.add_page_break()
        
        # 4. DANH MỤC CÁC CHỮ VIẾT TẮT VÀ KÝ HIỆU (Chuẩn VN: 2 cột)
        if abbreviations and len(abbreviations) > 0:
            doc.add_paragraph("DANH MỤC CÁC CHỮ VIẾT TẮT VÀ KÝ HIỆU", style='Front Heading')
            
            # Separate abbreviations and symbols
            abbr_list = [a for a in abbreviations if a.get('type') == 'abbreviation']
            symbol_list = [a for a in abbreviations if a.get('type') == 'symbol']
            
            # Sort alphabetically
            abbr_list.sort(key=lambda x: x.get('abbreviation', ''))
            symbol_list.sort(key=lambda x: x.get('abbreviation', ''))
            
            # Create table for abbreviations (2 columns - VN standard)
            if abbr_list:
                # Sub-heading for abbreviations
                p = doc.add_paragraph()
                run = p.add_run("Chữ viết tắt")
                set_font_complex(run.font, settings.font_family, settings.font_size, bold=True)
                
                # Table with header row - 2 columns
                abbr_table = doc.add_table(rows=len(abbr_list) + 1, cols=2)
                abbr_table.style = 'Table Grid'
                
                # Header row
                header_row = abbr_table.rows[0]
                headers = ["Chữ viết tắt", "Diễn giải đầy đủ"]
                for col_idx, header_text in enumerate(headers):
                    cell = header_row.cells[col_idx]
                    cell.text = header_text
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            set_font_complex(run.font, settings.font_family, settings.font_size, bold=True)
                
                # Data rows
                for idx, item in enumerate(abbr_list):
                    row = abbr_table.rows[idx + 1]  # +1 to skip header
                    # Abbreviation column (center)
                    cell0 = row.cells[0]
                    cell0.text = item.get('abbreviation', '')
                    for para in cell0.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            set_font_complex(run.font, settings.font_family, settings.font_size)
                    # Full form column (left)
                    cell1 = row.cells[1]
                    cell1.text = item.get('fullForm', '')
                    for para in cell1.paragraphs:
                        for run in para.runs:
                            set_font_complex(run.font, settings.font_family, settings.font_size)
                
                doc.add_paragraph()  # Spacing
            
            # Create table for symbols (2 columns - VN standard)
            if symbol_list:
                # Sub-heading for symbols
                p = doc.add_paragraph()
                run = p.add_run("Ký hiệu")
                set_font_complex(run.font, settings.font_family, settings.font_size, bold=True)
                
                # Table with header row - 2 columns
                symbol_table = doc.add_table(rows=len(symbol_list) + 1, cols=2)
                symbol_table.style = 'Table Grid'
                
                # Header row
                header_row = symbol_table.rows[0]
                headers = ["Ký hiệu", "Diễn giải đầy đủ"]
                for col_idx, header_text in enumerate(headers):
                    cell = header_row.cells[col_idx]
                    cell.text = header_text
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            set_font_complex(run.font, settings.font_family, settings.font_size, bold=True)
                
                # Data rows
                for idx, item in enumerate(symbol_list):
                    row = symbol_table.rows[idx + 1]  # +1 to skip header
                    # Symbol column (center, italic)
                    cell0 = row.cells[0]
                    cell0.text = item.get('abbreviation', '')
                    for para in cell0.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            set_font_complex(run.font, settings.font_family, settings.font_size, italic=True)
                    # Full form column (left)
                    cell1 = row.cells[1]
                    cell1.text = item.get('fullForm', '')
                    for para in cell1.paragraphs:
                        for run in para.runs:
                            set_font_complex(run.font, settings.font_family, settings.font_size)
            
            doc.add_page_break()
        
        # --- CONTENT ---
        h1_count = 0
        h2_count = 0
        h3_count = 0
        h4_count = 0
        h5_count = 0
        
        lines = text.split("\n")
        i = 0
        pending_table_caption = None

        while i < len(lines):
            line = lines[i]
            stripped_line = line.strip()
            
            # Table detection
            if stripped_line.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                
                # Process table
                if len(table_lines) >= 2:
                    rows_data = []
                    for tl in table_lines:
                        if "---" in tl: continue
                        cells = [c.strip() for c in tl.split("|")[1:-1]]
                        rows_data.append(cells)
                    
                    if rows_data:
                        # Add spacing before table
                        doc.add_paragraph() 

                        rows = len(rows_data)
                        cols = len(rows_data[0])
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        
                        for r in range(rows):
                            row_cells = table.rows[r].cells
                            for c in range(min(cols, len(rows_data[r]))):
                                cell = row_cells[c]
                                cell.text = rows_data[r][c]
                                for paragraph in cell.paragraphs:
                                    set_font_complex(paragraph.runs[0].font, settings.font_family, settings.font_size, 
                                                     bold=(r==0), color=RGBColor(0, 0, 0))
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Render Pending Caption BELOW Table
                        if pending_table_caption:
                            # Use Table Caption Style
                            doc.add_paragraph(pending_table_caption, style='Table Caption')
                            pending_table_caption = None
                        
                        # Add spacing after table group
                        doc.add_paragraph()

                continue 
            
            # Flush pending caption if not followed by a table immediately
            if pending_table_caption:
                doc.add_paragraph(pending_table_caption, style='Table Caption')
                pending_table_caption = None

            # Check for Table Captions: Bảng 1.1: Caption
            if stripped_line.startswith("Bảng") and ":" in stripped_line:
                pending_table_caption = stripped_line
                i += 1
                continue

            # Heading 1: CHƯƠNG I
            if line.startswith("# "):
                h1_count += 1
                h2_count = 0
                h3_count = 0
                h4_count = 0
                from core.utils.helpers import to_roman
                roman = to_roman(h1_count)
                raw_title = line[2:].strip()
                
                if settings.auto_numbering:
                    # Logic tách dòng tiêu đề chương
                    if settings.h1_split:
                        # Sử dụng \n để tạo soft break (Shift+Enter) trong Word
                        # Điều này giúp tiêu đề vẫn là 1 paragraph (tốt cho TOC) nhưng hiển thị 2 dòng
                        prefix = f"{settings.h1_prefix} {roman}" if settings.h1_prefix else f"{roman}"
                        full_title = f"{prefix}\n{raw_title.upper()}"
                    else:
                        prefix = f"{settings.h1_prefix} {roman}: " if settings.h1_prefix else f"{roman}. "
                        full_title = f"{prefix}{raw_title.upper()}"
                else:
                    full_title = raw_title.upper()
                
                heading = doc.add_paragraph(full_title, style='Heading 1')
                for run in heading.runs:
                    set_font_complex(run.font, settings.font_family, settings.h1_size, bold=True, color=RGBColor(0, 0, 0))
            
            # Heading 2: 1.1
            elif line.startswith("## "):
                h2_count += 1
                h3_count = 0
                h4_count = 0
                raw_title = line[3:].strip()
                if settings.auto_numbering:
                    # Nếu hierarchical_numbering = True -> 1.1, ngược lại -> 1
                    prefix = f"{h1_count}.{h2_count}." if settings.hierarchical_numbering else f"{h2_count}."
                    full_title = f"{prefix} {raw_title}"
                else:
                    full_title = raw_title
                
                heading = doc.add_paragraph(full_title, style='Heading 2')
                for run in heading.runs:
                    set_font_complex(run.font, settings.font_family, settings.h2_size, bold=True, color=RGBColor(0, 0, 0))
            
            # Heading 3: 1.1.1
            elif line.startswith("### "):
                h3_count += 1
                h4_count = 0
                raw_title = line[4:].strip()
                if settings.auto_numbering:
                    prefix = f"{h1_count}.{h2_count}.{h3_count}." if settings.hierarchical_numbering else f"{h2_count}.{h3_count}."
                    full_title = f"{prefix} {raw_title}"
                else:
                    full_title = raw_title
                
                heading = doc.add_paragraph(full_title, style='Heading 3')
                for run in heading.runs:
                    set_font_complex(run.font, settings.font_family, settings.h3_size, bold=True, italic=True, color=RGBColor(0, 0, 0))

            # Heading 4
            elif line.startswith("#### "):
                h4_count += 1
                h5_count = 0 # Reset h5_count
                raw_title = line[5:].strip()
                
                if settings.auto_numbering:
                    prefix = f"{h1_count}.{h2_count}.{h3_count}.{h4_count}." if settings.hierarchical_numbering else f"{h2_count}.{h3_count}.{h4_count}."
                    full_title = f"{prefix} {raw_title}"
                else:
                    full_title = raw_title
                
                # Word usually doesn't have Heading 4 by default or it's small. We'll use Heading 4 style if exists or Normal bold italic
                try:
                    heading = doc.add_paragraph(full_title, style='Heading 4')
                except:
                    heading = doc.add_paragraph(full_title, style='Normal')
                
                for run in heading.runs:
                    set_font_complex(run.font, settings.font_family, settings.font_size, bold=True, italic=True, color=RGBColor(0, 0, 0))
                    set_font_complex(run.font, settings.font_family, settings.font_size, bold=True, italic=True, color=RGBColor(0, 0, 0))
            
            # Heading 5
            elif line.startswith("##### "):
                h5_count += 1
                raw_title = line[6:].strip()
                
                if settings.auto_numbering:
                    prefix = f"{h1_count}.{h2_count}.{h3_count}.{h4_count}.{h5_count}." if settings.hierarchical_numbering else f"{h2_count}.{h3_count}.{h4_count}.{h5_count}."
                    full_title = f"{prefix} {raw_title}"
                else:
                    full_title = raw_title
                
                try:
                    heading = doc.add_paragraph(full_title, style='Heading 5')
                except:
                    heading = doc.add_paragraph(full_title, style='Normal')
                
                for run in heading.runs:
                    set_font_complex(run.font, settings.font_family, settings.font_size, italic=True, color=RGBColor(0, 0, 0))

            # Bullet points: - , -- , ---, or •, ●
            elif re.match(r'^(\-{1,3}|\*{1,3}|•|●)\s', line):
                match = re.match(r'^(\-{1,3}|\*{1,3}|•|●)\s', line)
                prefix = match.group(1)
                level = len(prefix)
                content = line[len(prefix)+1:].strip()
                
                # Chuẩn đồ án Việt Nam: dùng gạch đầu dòng (–) và dấu cộng (+)
                # Sử dụng En Dash (–) thay vì Hyphen (-) cho đẹp hơn
                bullet_char = "\u2013" # En Dash
                print(f"DEBUG: Processing bullet level {level}, char: {bullet_char}")
                
                if level == 2:
                    bullet_char = "+"
                elif level == 3:
                    bullet_char = "\u2013"
                
                # Manual bullet implementation using Normal style + Indent
                p = doc.add_paragraph(style='Normal')
                
                # Calculate indent
                # Base indent (e.g. 1.27cm) + Level indent
                base_indent = settings.indent # cm
                level_indent = (level - 1) * 0.75 # cm
                total_indent = base_indent + level_indent
                
                # Hanging indent logic:
                # Left Indent = Total Indent + Hanging Amount
                # First Line Indent = -Hanging Amount
                # This makes the bullet sit at Total Indent, and text wrap nicely.
                hanging = 0.5 # cm (reduced from 0.63 for tighter look)
                
                p_fmt = p.paragraph_format
                p_fmt.left_indent = Cm(total_indent + hanging)
                p_fmt.first_line_indent = Cm(-hanging)
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_fmt.tab_stops.add_tab_stop(Cm(total_indent + hanging))
                
                # Add bullet char and tab
                run = p.add_run(f"{bullet_char}\t")
                set_font_complex(run.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
                
                # Add content
                # We need to process formatting in content as well (bold, italic, latex)
                # Reuse the logic? Or just simple add for now?
                # Let's reuse the simple formatting logic by splitting
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|<u>.*?</u>|\$.*?\$)', content)
                for part in parts:
                    if not part: continue
                     # LaTeX handling
                    if part.startswith('$') and part.endswith('$'):
                        latex_content = part
                        # Handle display math inside bullet? usually inline
                        inner_tex = part
                        if part.startswith('$$'): inner_tex = part[2:-2]
                        elif part.startswith('$'): inner_tex = part[1:-1]
                        
                        image_stream, h_in, w_in, descent_in = render_latex_to_image(inner_tex, font_size_pt=settings.font_size)
                        if image_stream:
                            run = p.add_run()
                            # Convert inches to docx Length (Cm or Inches)
                            run.add_picture(image_stream, height=Cm(h_in * 2.54))
                            # Lower the image by descent amount to align baseline
                            # descent_in is in inches. 1 inch = 72 points.
                            set_run_position(run, -descent_in * 72)
                        else:
                            run = p.add_run(part)
                            set_font_complex(run.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
                        continue

                    run = p.add_run()
                    text_content = part
                    is_bold = False
                    is_italic = False
                    is_underline = False
                    
                    if part.startswith('**') and part.endswith('**'):
                        text_content = part[2:-2]
                        is_bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        text_content = part[1:-1]
                        is_italic = True
                    elif part.startswith('<u>') and part.endswith('</u>'):
                        text_content = part[3:-4]
                        is_underline = True
                        
                    run.text = text_content
                    set_font_complex(run.font, settings.font_family, settings.font_size, 
                                     bold=is_bold, italic=is_italic, color=RGBColor(0, 0, 0))
                    if is_underline:
                        run.font.underline = True
            
            elif line.strip():
                stripped_line = line.strip()
                print(f"DEBUG LINE: '{stripped_line}'")
                
                # Check for Figure placeholders: [Hình 1.1: Caption]
                # Relaxed regex to handle potential spaces
                match = re.match(r'\[\s*(Hình \d+\.\d+)\s*:\s*(.*)\s*\]', stripped_line)
                if match:
                    fig_num = match.group(1)
                    caption = match.group(2)
                    print(f"DEBUG: Found placeholder for {fig_num}")
                    
                    # Find figure data
                    # Normalize spaces for comparison
                    fig_data = next((f for f in figures if f.number.replace(" ", "") == fig_num.replace(" ", "")), None)
                    
                    if fig_data:
                        print(f"DEBUG: Found figure data: {fig_data.path}")
                        if fig_data.path:
                            try:
                                # Resolve path if relative
                                img_path = fig_data.path
                                if not os.path.isabs(img_path):
                                    img_path = os.path.abspath(img_path)
                                
                                print(f"DEBUG: Checking image path: {img_path}")
                                if os.path.exists(img_path):
                                    # Add image
                                    # Use width from figure data if available, else default to 16cm
                                    width = Cm(fig_data.width) if hasattr(fig_data, 'width') and fig_data.width else Cm(16)
                                    
                                    # Create a new paragraph for the image
                                    p = doc.add_paragraph()
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Reset indentation to ensure true center
                                    p.paragraph_format.left_indent = Cm(0)
                                    p.paragraph_format.right_indent = Cm(0)
                                    p.paragraph_format.first_line_indent = Cm(0)
                                    
                                    run = p.add_run()
                                    run.add_picture(img_path, width=width)
                                    print(f"DEBUG: Inserted image {img_path}")
                                    
                                    # Add caption
                                    caption_text = f"{fig_num}: {caption}"
                                    # Use 'Figure Caption' style which is defined as Black/Italic
                                    caption_p = doc.add_paragraph(caption_text, style='Figure Caption')
                                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                else:
                                    print(f"DEBUG: Image path not found: {img_path}")
                                    doc.add_paragraph(f"[Hình ảnh không tìm thấy: {img_path}]", style='Normal')
                            except Exception as e:
                                print(f"Error adding image {img_path}: {e}")
                                doc.add_paragraph(f"[Lỗi chèn hình: {fig_num}]", style='Normal')
                        else:
                             print("DEBUG: Figure path is empty")
                    else:
                        print(f"DEBUG: Figure data not found for {fig_num}")
                        doc.add_paragraph(f"[Hình ảnh không tìm thấy dữ liệu: {fig_num}]", style='Normal')
                
                else:
                    # Normal paragraph with formatting support
                    p = doc.add_paragraph(style='Normal')
                    
                    # Regex to split by formatting tokens: **bold**, *italic*, <u>underline</u>, $$display$$, $inline$
                    # Priority: $$ first
                    parts = re.split(r'(\$\$.*?\$\$|\$.*?\$|\*\*.*?\*\*|\*.*?\*|<u>.*?</u>)', stripped_line)
                    
                    for part in parts:
                        if not part: continue
                        
                        # LaTeX handling
                        if part.startswith('$') and part.endswith('$'):
                            # Handle display math $$...$$ vs inline $...$
                            is_display = part.startswith('$$')
                            if is_display:
                                latex_content = part[2:-2]
                            else:
                                latex_content = part[1:-1]
                            
                            # Try native OMML first (best quality)
                            if insert_omml_equation(p, latex_content):
                                continue
                            
                            # Fallback to image rendering
                            image_stream, h_in, w_in, descent_in = render_latex_to_image(latex_content, font_size_pt=settings.font_size, is_display=is_display)
                            if image_stream:
                                scaled_height = h_in * 2.54
                                if is_display:
                                    # Display math: create new centered paragraph
                                    math_p = doc.add_paragraph()
                                    math_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    math_p.paragraph_format.first_line_indent = Cm(0)
                                    math_p.paragraph_format.space_before = Pt(6)
                                    math_p.paragraph_format.space_after = Pt(6)
                                    run = math_p.add_run()
                                    run.add_picture(image_stream, height=Cm(scaled_height))
                                    # Create new paragraph for remaining text
                                    p = doc.add_paragraph(style='Normal')
                                else:
                                    run = p.add_run()
                                    run.add_picture(image_stream, height=Cm(scaled_height))
                                    # Adjust baseline shift proportionally
                                    set_run_position(run, -descent_in * 72)
                            else:
                                # Fallback: Show placeholder for complex formulas
                                if is_display:
                                    # Create centered paragraph with placeholder
                                    math_p = doc.add_paragraph()
                                    math_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    math_p.paragraph_format.first_line_indent = Cm(0)
                                    math_p.paragraph_format.space_before = Pt(6)
                                    math_p.paragraph_format.space_after = Pt(6)
                                    run = math_p.add_run(f"[Công thức phức tạp - cần chèn thủ công]")
                                    set_font_complex(run.font, settings.font_family, settings.font_size, italic=True, color=RGBColor(128, 128, 128))
                                    # Create new paragraph for remaining text
                                    p = doc.add_paragraph(style='Normal')
                                else:
                                    # Inline: just show simple placeholder
                                    run = p.add_run("[công thức]")
                                    set_font_complex(run.font, settings.font_family, settings.font_size, italic=True, color=RGBColor(128, 128, 128))
                            continue

                        run = p.add_run()
                        text_content = part
                        is_bold = False
                        is_italic = False
                        is_underline = False
                        
                        if part.startswith('**') and part.endswith('**'):
                            text_content = part[2:-2]
                            is_bold = True
                        elif part.startswith('*') and part.endswith('*'):
                            text_content = part[1:-1]
                            is_italic = True
                        elif part.startswith('<u>') and part.endswith('</u>'):
                            text_content = part[3:-4]
                            is_underline = True
                            
                        run.text = text_content
                        set_font_complex(run.font, settings.font_family, settings.font_size, 
                                         bold=is_bold, italic=is_italic, color=RGBColor(0, 0, 0))
                        if is_underline:
                            run.font.underline = True
                    
                    p.paragraph_format.first_line_indent = Cm(settings.indent)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        set_font_complex(run.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
            
            i += 1
                
        # --- REFERENCES ---
        if citations:
            doc.add_page_break()
            ref_p = doc.add_paragraph("TÀI LIỆU THAM KHẢO", style='Front Heading')
            for i, c in enumerate(citations, 1):
                p = doc.add_paragraph(f"[{i}] {format_citation_apa(c)}")
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                     set_font_complex(run.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
        
        doc.save(file_path)
        return True, f"Đã xuất file Word:\n{file_path}"
        
    except Exception as e:
        traceback.print_exc()
        return False, f"Không thể xuất file Word:\n{str(e)}"
