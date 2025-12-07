from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ThesisFormatter.models.data_classes import Settings, Figure, Table, Citation
from ThesisFormatter.utils.helpers import format_citation_apa
from typing import List
import os
import re
import traceback

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_toc_field(paragraph, field_code):
    """Insert a Word field (like TOC) into a paragraph"""
    run = paragraph.add_run()
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'begin')
    run._element.append(fldChar)

    run = paragraph.add_run()
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = field_code
    run._element.append(instrText)

    run = paragraph.add_run()
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'separate')
    run._element.append(fldChar)

    run = paragraph.add_run()
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'end')
    run._element.append(fldChar)

def set_font_complex(font, font_name, font_size=None, bold=False, italic=False, color=None):
    """Set font properties comprehensively for all script types"""
    font.name = font_name
    if font_size:
        font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color
        
    # Access the underlying XML element to set fonts for all regions
    rPr = font._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def setup_styles(doc, settings: Settings):
    """Configure document styles to match thesis requirements"""
    styles = doc.styles
    
    # Normal Style
    style_normal = styles['Normal']
    set_font_complex(style_normal.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
    paragraph_format = style_normal.paragraph_format
    paragraph_format.line_spacing = settings.line_spacing
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Heading 1
    style_h1 = styles['Heading 1']
    set_font_complex(style_h1.font, settings.font_family, settings.h1_size, bold=True, color=RGBColor(0, 0, 0))
    paragraph_format = style_h1.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_before = Pt(24)
    paragraph_format.space_after = Pt(18)
    
    # Heading 2
    style_h2 = styles['Heading 2']
    set_font_complex(style_h2.font, settings.font_family, settings.h2_size, bold=True, color=RGBColor(0, 0, 0))
    paragraph_format = style_h2.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format.space_before = Pt(18)
    paragraph_format.space_after = Pt(12)
    
    # Heading 3
    style_h3 = styles['Heading 3']
    set_font_complex(style_h3.font, settings.font_family, settings.h3_size, bold=True, italic=True, color=RGBColor(0, 0, 0))
    paragraph_format = style_h3.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(6)
    
    # Front Heading (for TOC, List of Tables, etc. - Looks like H1 but not in TOC)
    try:
        style_front = styles.add_style('Front Heading', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style_front = styles['Front Heading']
    style_front.base_style = styles['Normal']
    set_font_complex(style_front.font, settings.font_family, settings.h1_size, bold=True, color=RGBColor(0, 0, 0))
    style_front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_front.paragraph_format.space_before = Pt(24)
    style_front.paragraph_format.space_after = Pt(18)

    # Figure Caption Style
    try:
        style_fig_cap = styles.add_style('Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style_fig_cap = styles['Figure Caption']
    style_fig_cap.base_style = styles['Normal']
    set_font_complex(style_fig_cap.font, settings.font_family, settings.font_size, italic=True, color=RGBColor(0, 0, 0))
    style_fig_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_fig_cap.paragraph_format.space_before = Pt(6)
    style_fig_cap.paragraph_format.space_after = Pt(12)

    # Table Caption Style
    try:
        style_tbl_cap = styles.add_style('Table Caption', WD_STYLE_TYPE.PARAGRAPH)
    except:
        style_tbl_cap = styles['Table Caption']
    style_tbl_cap.base_style = styles['Normal']
    set_font_complex(style_tbl_cap.font, settings.font_family, settings.font_size, bold=True, color=RGBColor(0, 0, 0))
    style_tbl_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_tbl_cap.paragraph_format.space_before = Pt(12)
    style_tbl_cap.paragraph_format.space_after = Pt(6)
    
    # List Bullet
    try:
        style_bullet = styles['List Bullet']
        set_font_complex(style_bullet.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
    except:
        pass

def export_to_docx(file_path: str, text: str, settings: Settings, 
                   figures: List[Figure], tables: List[Table], citations: List[Citation]):
    try:
        doc = Document()
        setup_styles(doc, settings)
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Cm(settings.margin_top)
            section.bottom_margin = Cm(settings.margin_bottom)
            section.left_margin = Cm(settings.margin_left)
            section.right_margin = Cm(settings.margin_right)
        
        # --- FRONT MATTER ---
        
        # 1. MỤC LỤC
        # Add instruction to update fields
        instr_p = doc.add_paragraph("Lưu ý: Nhấn Ctrl+A rồi nhấn F9 (hoặc chuột phải chọn 'Update Field') để cập nhật Mục lục và Danh mục.")
        instr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_complex(instr_p.runs[0].font, settings.font_family, 11, italic=True, color=RGBColor(255, 0, 0))

        toc_p = doc.add_paragraph("MỤC LỤC", style='Front Heading')
        # Insert TOC Field: \o "1-3" includes Heading 1-3, \h hyperlinks, \z hide page numbers in web, \u outline levels
        p = doc.add_paragraph()
        add_toc_field(p, r'TOC \o "1-3" \h \z \u')
        doc.add_page_break()
        
        # 2. DANH MỤC HÌNH ẢNH
        if figures:
            doc.add_paragraph("DANH MỤC HÌNH ẢNH", style='Front Heading')
            # Insert TOC for Figure Caption style
            p = doc.add_paragraph()
            add_toc_field(p, r'TOC \h \z \t "Figure Caption,1"')
            doc.add_page_break()
            
        # 3. DANH MỤC BẢNG BIỂU
        # Note: We don't have a tables list passed in explicitly with captions usually, 
        # but we detect them in text. However, for the TOC to work, we just need the captions in text to use the style.
        # We'll assume if there are tables in text, we want this list. 
        # Since we don't track tables list perfectly in backend yet, we'll just add it if text contains "Bảng"
        if "Bảng" in text:
            doc.add_paragraph("DANH MỤC BẢNG BIỂU", style='Front Heading')
            p = doc.add_paragraph()
            add_toc_field(p, r'TOC \h \z \t "Table Caption,1"')
            doc.add_page_break()
        
        # --- CONTENT ---
        h1_count = 0
        h2_count = 0
        h3_count = 0
        
        lines = text.split("\n")
        i = 0
        pending_table_caption = None

        while i < len(lines):
            line = lines[i]
            stripped_line = line.strip()
            
            # Table detection
            if stripped_line.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                
                # Process table
                if len(table_lines) >= 2:
                    rows_data = []
                    for tl in table_lines:
                        if "---" in tl: continue
                        cells = [c.strip() for c in tl.split("|")[1:-1]]
                        rows_data.append(cells)
                    
                    if rows_data:
                        # Add spacing before table
                        doc.add_paragraph() 

                        rows = len(rows_data)
                        cols = len(rows_data[0])
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        
                        for r in range(rows):
                            row_cells = table.rows[r].cells
                            for c in range(min(cols, len(rows_data[r]))):
                                cell = row_cells[c]
                                cell.text = rows_data[r][c]
                                for paragraph in cell.paragraphs:
                                    set_font_complex(paragraph.runs[0].font, settings.font_family, settings.font_size, 
                                                     bold=(r==0), color=RGBColor(0, 0, 0))
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Render Pending Caption BELOW Table
                        if pending_table_caption:
                            # Use Table Caption Style
                            doc.add_paragraph(pending_table_caption, style='Table Caption')
                            pending_table_caption = None
                        
                        # Add spacing after table group
                        doc.add_paragraph()

                continue 
            
            # Flush pending caption if not followed by a table immediately
            if pending_table_caption:
                doc.add_paragraph(pending_table_caption, style='Table Caption')
                pending_table_caption = None

            # Check for Table Captions: Bảng 1.1: Caption
            if stripped_line.startswith("Bảng") and ":" in stripped_line:
                pending_table_caption = stripped_line
                i += 1
                continue

            # Heading 1: CHƯƠNG I
            if line.startswith("# "):
                h1_count += 1
                h2_count = 0
                h3_count = 0
                from ThesisFormatter.utils.helpers import to_roman
                roman = to_roman(h1_count)
                raw_title = line[2:].strip()
                full_title = f"CHƯƠNG {roman}: {raw_title.upper()}"
                
                heading = doc.add_paragraph(full_title, style='Heading 1')
                for run in heading.runs:
                    set_font_complex(run.font, settings.font_family, settings.h1_size, bold=True, color=RGBColor(0, 0, 0))
            
            # Heading 2: 1.1
            elif line.startswith("## "):
                h2_count += 1
                h3_count = 0
                raw_title = line[3:].strip()
                full_title = f"{h1_count}.{h2_count}. {raw_title}"
                
                heading = doc.add_paragraph(full_title, style='Heading 2')
                for run in heading.runs:
                    set_font_complex(run.font, settings.font_family, settings.h2_size, bold=True, color=RGBColor(0, 0, 0))
            
            # Heading 3: 1.1.1
            elif line.startswith("### "):
                h3_count += 1
                raw_title = line[4:].strip()
                full_title = f"{h1_count}.{h2_count}.{h3_count}. {raw_title}"
                
                heading = doc.add_paragraph(full_title, style='Heading 3')
                for run in heading.runs:
                    set_font_complex(run.font, settings.font_family, settings.h3_size, bold=True, italic=True, color=RGBColor(0, 0, 0))
            
            elif line.startswith("- "):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                for run in p.runs:
                    set_font_complex(run.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
            
            elif line.strip():
                stripped_line = line.strip()
                # Check for Figure placeholders
                if stripped_line.startswith("[") and stripped_line.endswith("]"):
                    print(f"DEBUG: Processing potential figure line: '{stripped_line}'")
                    match = re.search(r"\[\s*(Hình\s+\d+\.\d+)\s*:", stripped_line)
                    if match:
                        fig_num = match.group(1)
                        figure = next((f for f in figures if f.number.replace(" ", "") == fig_num.replace(" ", "")), None)
                        
                        if figure:
                            img_path = figure.path
                            if not os.path.isabs(img_path):
                                img_path = os.path.abspath(img_path)
                                
                            if os.path.exists(img_path):
                                try:
                                    # Add spacing before image
                                    doc.add_paragraph()
                                    
                                    run = doc.add_paragraph().add_run()
                                    width_cm = figure.width if hasattr(figure, 'width') and figure.width else 16
                                    print(f"DEBUG: Inserting image {img_path} with width {width_cm} cm")
                                    run.add_picture(img_path, width=Cm(width_cm))
                                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    print(f"DEBUG: Successfully inserted image from {img_path}")
                                except Exception as e:
                                    print(f"ERROR adding picture to docx: {e}")
                                    traceback.print_exc()
                            else:
                                print(f"ERROR: Image file does not exist at: {img_path}")
                    
                    # Add Caption BELOW image using Figure Caption Style
                    caption_text = stripped_line.strip("[]")
                    doc.add_paragraph(caption_text, style='Figure Caption')
                    
                    # Add spacing after image group
                    doc.add_paragraph()
                    
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.first_line_indent = Cm(settings.indent)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        set_font_complex(run.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
            
            i += 1
                
        # --- REFERENCES ---
        if citations:
            doc.add_page_break()
            ref_p = doc.add_paragraph("TÀI LIỆU THAM KHẢO", style='Front Heading')
            for i, c in enumerate(citations, 1):
                p = doc.add_paragraph(f"[{i}] {format_citation_apa(c)}")
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                     set_font_complex(run.font, settings.font_family, settings.font_size, color=RGBColor(0, 0, 0))
        
        doc.save(file_path)
        return True, f"Đã xuất file Word:\n{file_path}"
        
    except Exception as e:
        traceback.print_exc()
        return False, f"Không thể xuất file Word:\n{str(e)}"
